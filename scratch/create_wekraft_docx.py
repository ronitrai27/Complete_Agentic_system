from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "data" / "wekraft_project_summary.docx"


def set_run(run, size: float, bold: bool = False, color: str = "1F2937") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_fact(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    set_run(paragraph.add_run(text), 11)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2563EB", 14, 7),
    ("Heading 2", 13, "2563EB", 10, 5),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("WEKRAFT | PROJECT SUMMARY"), 8.5, bold=True, color="64748B")

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
set_run(title.add_run("Project WeKraft"), 25, bold=True, color="0F172A")

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
set_run(
    subtitle.add_run("A compact knowledge-graph ingestion test"),
    12,
    color="475569",
)

metadata = doc.add_paragraph()
metadata.paragraph_format.space_after = Pt(12)
set_run(metadata.add_run("Organization: "), 10, bold=True, color="2563EB")
set_run(metadata.add_run("NovaSphere Technologies.  "), 10, color="0F172A")
set_run(metadata.add_run("Status: "), 10, bold=True, color="2563EB")
set_run(metadata.add_run("Active."), 10, color="0F172A")

doc.add_heading("Project overview.", level=1)
overview = doc.add_paragraph()
set_run(
    overview.add_run(
        "Project WeKraft is a collaboration product at NovaSphere Technologies."
    ),
    11,
)

doc.add_heading("Team and connected work.", level=1)
facts = [
    "Ronit Rai leads Project WeKraft.",
    "Akash Sharma develops the WeKraft API.",
    "Riya Kapoor designs the WeKraft user experience.",
    "Mia Chen tests the WeKraft release.",
    "Project WeKraft depends on AWS infrastructure.",
    "Project WeKraft follows Security Policy v14.",
    "Project WeKraft collaborates with Engineering Department.",
    "Project WeKraft extends Project Atlas.",
    "Project WeKraft supports Project Orion.",
    "Ronit Rai reports to Product Department.",
    "Akash Sharma works in Security Department.",
    "Riya Kapoor collaborates with Customer Success Department.",
    "Mia Chen collaborates with Research Department.",
]
for fact in facts:
    add_fact(doc, fact)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(
    footer.add_run("Prepared for AI Flow ingestion and relationship validation"),
    8.5,
    color="64748B",
)

doc.core_properties.title = "Project WeKraft Summary"
doc.core_properties.subject = "Knowledge graph ingestion test"
doc.core_properties.author = "AI Flow"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
